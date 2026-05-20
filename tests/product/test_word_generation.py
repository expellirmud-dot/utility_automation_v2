"""
Tests for PRODUCT-004: Word Memo / Dika Document Generation.

Covers:
- thai_date_format helper
- bahttext helper: exact amounts (11, 21, 101, 3210, 1250.50, 0, 1M)
- WordGenerationService: output file created, template unchanged, placeholder replaced
- API: POST /dika, POST /memo/generate, GET /memo/download
"""

import datetime
import os
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest
from fastapi.testclient import TestClient

from src.product.services.word_generation import thai_date_format, bahttext, WordGenerationService
from src.product.main import app

client = TestClient(app)

TEMPLATE_PATH = Path("Word Template.docx")
OUTPUT_DIR = Path("data/generated")


# ─────────────────────────────────────────────
# thai_date_format
# ─────────────────────────────────────────────

class TestThaiDateFormat:
    def test_normal_date(self):
        assert thai_date_format(datetime.date(2026, 5, 20)) == "20 พฤษภาคม 2569"

    def test_january(self):
        assert thai_date_format(datetime.date(2025, 1, 1)) == "1 มกราคม 2568"

    def test_none_returns_empty(self):
        assert thai_date_format(None) == ""

    def test_october_fiscal_month(self):
        assert thai_date_format(datetime.date(2025, 10, 15)) == "15 ตุลาคม 2568"

    def test_december(self):
        assert thai_date_format(datetime.date(2025, 12, 31)) == "31 ธันวาคม 2568"


# ─────────────────────────────────────────────
# bahttext — exact amounts
# ─────────────────────────────────────────────

class TestBahttext:
    def test_zero(self):
        assert bahttext(0.0) == "ศูนย์บาทถ้วน"

    def test_eleven(self):
        # 11 -> สิบหนึ่งบาทถ้วน
        result = bahttext(11.0)
        assert result == "สิบหนึ่งบาทถ้วน"

    def test_twenty_one(self):
        # 21 -> ยี่สิบหนึ่งบาทถ้วน
        result = bahttext(21.0)
        assert result == "ยี่สิบหนึ่งบาทถ้วน"

    def test_one_hundred_one(self):
        # 101 -> หนึ่งร้อยหนึ่งบาทถ้วน
        result = bahttext(101.0)
        assert result == "หนึ่งร้อยหนึ่งบาทถ้วน"

    def test_three_thousand_two_hundred_ten(self):
        # 3210 -> สามพันสองร้อยสิบบาทถ้วน
        result = bahttext(3210.0)
        assert result == "สามพันสองร้อยสิบบาทถ้วน"

    def test_1250_50_satang(self):
        # 1250.50 -> หนึ่งพันสองร้อยห้าสิบบาทห้าสิบสตางค์
        result = bahttext(1250.50)
        assert result == "หนึ่งพันสองร้อยห้าสิบบาทห้าสิบสตางค์"

    def test_one_million_exact(self):
        result = bahttext(1_000_000.0)
        assert "ล้าน" in result
        assert "ถ้วน" in result

    def test_one_baht_fifty_satang(self):
        result = bahttext(1.50)
        assert "หนึ่งบาท" in result
        assert "ห้าสิบสตางค์" in result

    def test_no_satang_ends_in_thuai(self):
        result = bahttext(500.0)
        assert result.endswith("ถ้วน"), f"Expected ถ้วน suffix, got: {result}"

    def test_satang_does_not_end_in_thuai(self):
        result = bahttext(500.25)
        assert "สตางค์" in result
        assert not result.endswith("ถ้วน")


# ─────────────────────────────────────────────
# WordGenerationService
# ─────────────────────────────────────────────

class TestWordGenerationService:
    def setup_method(self):
        test_output = OUTPUT_DIR / "case_99999_memo.docx"
        if test_output.exists():
            test_output.unlink()

    def teardown_method(self):
        test_output = OUTPUT_DIR / "case_99999_memo.docx"
        if test_output.exists():
            test_output.unlink()

    def _generate(self, **kwargs):
        defaults = dict(
            case_id=99999,
            memo_number="TEST/001",
            memo_date=datetime.date(2026, 5, 20),
            expense_group="ค่าบริการโทรศัพท์",
            work_month="พฤษภาคม",
            fiscal_year_be=2569,
            provider="NT",
            bill_date=datetime.date(2026, 5, 10),
            total_amount=1250.0,
        )
        defaults.update(kwargs)
        return WordGenerationService.generate(**defaults)

    def test_template_exists(self):
        assert TEMPLATE_PATH.exists(), "Word Template.docx must exist in project root"

    def test_template_not_modified(self):
        if not TEMPLATE_PATH.exists():
            pytest.skip("Template not available")
        original_size = TEMPLATE_PATH.stat().st_size
        original_mtime = TEMPLATE_PATH.stat().st_mtime
        self._generate()
        assert TEMPLATE_PATH.stat().st_size == original_size, "Template file size changed!"
        assert TEMPLATE_PATH.stat().st_mtime == original_mtime, "Template mtime changed — template was modified!"

    def test_output_file_created(self):
        if not TEMPLATE_PATH.exists():
            pytest.skip("Template not available")
        output = self._generate()
        assert output.exists(), f"Generated file not found: {output}"
        assert output.stat().st_size > 0, "Generated file is empty"

    def test_fiscal_year_placeholder_replaced(self):
        if not TEMPLATE_PATH.exists():
            pytest.skip("Template not available")
        from docx import Document
        output = self._generate(fiscal_year_be=2569)
        doc = Document(output)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[…25xx..]" not in full_text, "Fiscal year placeholder was not replaced"
        assert "2569" in full_text

    def test_expense_group_placeholder_replaced(self):
        if not TEMPLATE_PATH.exists():
            pytest.skip("Template not available")
        from docx import Document
        output = self._generate(expense_group="ค่าน้ำประปา")
        doc = Document(output)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[ ค่า....................]" not in full_text
        assert "ค่าน้ำประปา" in full_text

    def test_memo_date_placeholder_replaced(self):
        if not TEMPLATE_PATH.exists():
            pytest.skip("Template not available")
        from docx import Document
        output = self._generate(memo_date=datetime.date(2026, 5, 20))
        doc = Document(output)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "[วันที่ทำรายการ]" not in full_text
        assert "พฤษภาคม" in full_text

    def test_missing_template_raises(self):
        import src.product.services.word_generation as svc
        original = svc.TEMPLATE_PATH
        try:
            svc.TEMPLATE_PATH = Path("nonexistent_template.docx")
            with pytest.raises(FileNotFoundError):
                WordGenerationService.generate(
                    case_id=1, memo_number="X", memo_date=None,
                    expense_group="test", work_month="มกราคม",
                    fiscal_year_be=2569, provider="",
                    bill_date=None, total_amount=0.0,
                )
        finally:
            svc.TEMPLATE_PATH = original


# ─────────────────────────────────────────────
# API tests
# ─────────────────────────────────────────────

from src.product.db.models import Base
from src.product.db.session import engine


@pytest.fixture(autouse=True)
def setup_database():
    """Create all tables before each test, drop after."""
    Base.metadata.create_all(bind=engine)
    yield
    Base.metadata.drop_all(bind=engine)


class TestDikaAndMemoAPI:
    """
    API integration tests for Dika save, memo generate, and download routes.
    DB is created fresh per test via the autouse fixture above.
    WordGenerationService.generate is mocked to avoid filesystem dependency.
    """

    def _create_case(self) -> int:
        resp = client.post("/api/cases/", json={
            "fiscal_year_be": 2569,
            "work_month": "พฤษภาคม",
            "case_type": "utility",
            "expense_group": "ค่าบริการโทรศัพท์",
            "department": "สำนักปลัดเทศบาล",
        })
        assert resp.status_code == 200, f"Case creation failed: {resp.text}"
        return resp.json()["id"]

    def test_save_dika_returns_200(self):
        case_id = self._create_case()
        resp = client.post(f"/api/cases/{case_id}/dika", json={
            "dika_no": "123/2569",
            "dika_date": "2026-05-20",
            "payee_name": "บริษัท โทรคมนาคมแห่งชาติ จำกัด (มหาชน)",
            "memo_number": "45",
            "memo_date": "2026-05-20",
        })
        assert resp.status_code == 200, f"Dika save failed: {resp.text}"
        data = resp.json()
        assert data["dika_number"] == "123/2569"
        assert data["payee_name"] == "บริษัท โทรคมนาคมแห่งชาติ จำกัด (มหาชน)"
        assert data["memo_number"] == "45"

    def test_save_dika_unknown_case_returns_404(self):
        resp = client.post("/api/cases/99999/dika", json={
            "dika_no": "X", "dika_date": "2026-05-20",
            "payee_name": "test", "memo_number": "1",
            "memo_date": "2026-05-20",
        })
        assert resp.status_code == 404

    def test_generate_memo_without_dika_returns_400(self):
        case_id = self._create_case()
        resp = client.post(f"/api/cases/{case_id}/memo/generate")
        assert resp.status_code == 400
        assert "Dika" in resp.json()["detail"]

    def test_generate_memo_returns_200(self):
        case_id = self._create_case()
        client.post(f"/api/cases/{case_id}/dika", json={
            "dika_no": "123/2569", "dika_date": "2026-05-20",
            "payee_name": "NT", "memo_number": "45",
            "memo_date": "2026-05-20",
        })
        fake_path = Path("data/generated") / f"case_{case_id}_memo.docx"
        with patch(
            "src.product.api.memos.WordGenerationService.generate",
            return_value=fake_path
        ):
            resp = client.post(f"/api/cases/{case_id}/memo/generate")
        assert resp.status_code == 200, f"Generate failed: {resp.text}"
        data = resp.json()
        assert data["status"] == "generated"
        assert f"/api/cases/{case_id}/memo/download" in data["download_url"]

    def test_download_memo_without_generated_file_returns_404(self):
        case_id = self._create_case()
        resp = client.get(f"/api/cases/{case_id}/memo/download")
        assert resp.status_code == 404
