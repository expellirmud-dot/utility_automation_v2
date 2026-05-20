"""
Word Memo Generation Service for PRODUCT-004.

Copies Word Template.docx (never modifies the original),
merges split runs within each paragraph, then performs
deterministic placeholder replacements using the context
mappings identified from the exact paragraph inspection.
"""

import shutil
import os
import datetime
from pathlib import Path
from typing import Optional

import docx
from docx import Document

# ─────────────────────────────────────────────
# Thai locale helpers
# ─────────────────────────────────────────────

_THAI_MONTHS = [
    "", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน",
    "พฤษภาคม", "มิถุนายน", "กรกฎาคม", "สิงหาคม",
    "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม",
]


def thai_date_format(d: Optional[datetime.date]) -> str:
    """Convert a date to Thai Buddhist Era date string.
    e.g. date(2026, 5, 20) -> '20 พฤษภาคม 2569'
    """
    if d is None:
        return ""
    be_year = d.year + 543
    return f"{d.day} {_THAI_MONTHS[d.month]} {be_year}"


_ONES = [
    "", "หนึ่ง", "สอง", "สาม", "สี่", "ห้า",
    "หก", "เจ็ด", "แปด", "เก้า",
]
_TENS = [
    "", "สิบ", "ยี่สิบ", "สามสิบ", "สี่สิบ", "ห้าสิบ",
    "หกสิบ", "เจ็ดสิบ", "แปดสิบ", "เก้าสิบ",
]
_SCALES = ["", "สิบ", "ร้อย", "พัน", "หมื่น", "แสน", "ล้าน"]


def _group_to_thai(n: int) -> str:
    """Convert integer 0-999999 to Thai word string."""
    if n == 0:
        return ""
    parts = []
    digits = []
    temp = n
    while temp > 0:
        digits.append(temp % 10)
        temp //= 10
    digits.reverse()
    length = len(digits)
    for i, d in enumerate(digits):
        pos = length - 1 - i
        if d == 0:
            continue
        if pos == 1 and d == 1:
            parts.append("สิบ")
        elif pos == 1 and d == 2:
            parts.append("ยี่สิบ")
        elif pos == 1:
            parts.append(_ONES[d] + "สิบ")
        elif pos == 0:
            parts.append(_ONES[d])
        else:
            parts.append(_ONES[d] + _SCALES[pos])
    return "".join(parts)


def bahttext(amount: float) -> str:
    """Convert float amount to Thai Baht text.
    e.g. 1250.50 -> 'หนึ่งพันสองร้อยห้าสิบบาทห้าสิบสตางค์'
         1250.00 -> 'หนึ่งพันสองร้อยห้าสิบบาทถ้วน'
    """
    if amount < 0:
        return "ลบ" + bahttext(-amount)

    total_satang = round(amount * 100)
    baht_int = total_satang // 100
    satang_int = total_satang % 100

    if baht_int == 0 and satang_int == 0:
        return "ศูนย์บาทถ้วน"

    result = ""
    if baht_int >= 1_000_000:
        millions = baht_int // 1_000_000
        remainder = baht_int % 1_000_000
        result += _group_to_thai(millions) + "ล้าน"
        if remainder > 0:
            result += _group_to_thai(remainder)
    else:
        result += _group_to_thai(baht_int)

    result += "บาท"

    if satang_int == 0:
        result += "ถ้วน"
    else:
        result += _group_to_thai(satang_int) + "สตางค์"

    return result


# ─────────────────────────────────────────────
# Paragraph-level run merger
# ─────────────────────────────────────────────

def _merge_runs(paragraph) -> None:
    """Merge all runs in a paragraph into the first run.

    Word XML frequently splits bracketed placeholder text across
    multiple runs due to formatting flags or cursor positions.
    Merging into a single run allows reliable str.replace() calls
    while preserving the paragraph-level formatting of the first run.
    """
    if not paragraph.runs:
        return
    full_text = "".join(r.text for r in paragraph.runs)
    # Write merged text into first run, clear the rest
    paragraph.runs[0].text = full_text
    for run in paragraph.runs[1:]:
        run.text = ""


# ─────────────────────────────────────────────
# Placeholder map builder
# ─────────────────────────────────────────────

def _build_replacements(
    memo_number: str,
    memo_date: Optional[datetime.date],
    expense_group: str,
    work_month: str,
    fiscal_year_be: int,
    provider: str,
    bill_date: Optional[datetime.date],
    total_amount: float,
) -> dict:
    """Return ordered list of (old, new) replacement pairs.

    Ordering matters: longer/more-specific strings are replaced first
    to prevent partial matches corrupting adjacent tokens.
    """
    date_str = thai_date_format(memo_date)
    bill_date_str = thai_date_format(bill_date)
    period_str = f"{work_month} {fiscal_year_be}"
    amount_str = f"{total_amount:,.2f}"
    amount_text = bahttext(total_amount)

    return [
        # Paragraph 5: memo reference number
        ("[….]",                              memo_number or ""),
        # Paragraph 7: memo date
        ("[วันที่ทำรายการ]",                  date_str),
        # Paragraphs 9 & 12: expense description (duplicate)
        ("[ ค่า....................]",         expense_group or ""),
        # Paragraph 12: billing period month+year
        ("[……….. 25xx…]",                    period_str),
        # Paragraph 12: provider/bill issuer
        ("[…………………ผู้ออกบิล…………………..]",   provider or ""),
        # Paragraph 12: invoice number placeholder — leave blank (no BillHeader field)
        ("[…………………..]",                     ""),
        # Paragraph 12: secondary bill reference — leave blank for single-bill MVP
        ("[……………(กรณีมากกว่า 1 บิล เช่นค่าโทรศัพท์สำนักงาน)………..] ", ""),
        # Paragraph 12: bill issue date
        ("[…….วันที่ออกบิล.......]",           bill_date_str),
        # Paragraph 12: total amount (numeric) — first occurrence
        # Paragraph 13: total amount (numeric) — duplicate; both resolved together
        ("[…………..]",                          amount_str),
        # Paragraph 12: total amount in Thai Baht words
        ("[……………….บาท…..สตางค์]",            amount_text),
        # Paragraph 13: fiscal year
        ("[…25xx..]",                          str(fiscal_year_be)),
        # Paragraph 13: remaining budget — stubbed
        ("[………………………….]",                   ""),
    ]


# ─────────────────────────────────────────────
# Main generation service
# ─────────────────────────────────────────────

TEMPLATE_PATH = Path("Word Template.docx")
OUTPUT_DIR = Path("data/generated")


class WordGenerationService:
    """Generates a filled Word memo from the canonical template."""

    @staticmethod
    def generate(
        case_id: int,
        memo_number: str,
        memo_date: Optional[datetime.date],
        expense_group: str,
        work_month: str,
        fiscal_year_be: int,
        provider: str,
        bill_date: Optional[datetime.date],
        total_amount: float,
    ) -> Path:
        """
        Copy template, fill placeholders, save output.

        Returns the absolute Path of the generated .docx file.
        Raises FileNotFoundError if template is missing.
        """
        if not TEMPLATE_PATH.exists():
            raise FileNotFoundError(f"Word template not found: {TEMPLATE_PATH}")

        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        output_path = OUTPUT_DIR / f"case_{case_id}_memo.docx"

        # Always copy fresh from template — never modify original
        shutil.copy2(TEMPLATE_PATH, output_path)

        doc = Document(output_path)

        replacements = _build_replacements(
            memo_number=memo_number,
            memo_date=memo_date,
            expense_group=expense_group,
            work_month=work_month,
            fiscal_year_be=fiscal_year_be,
            provider=provider,
            bill_date=bill_date,
            total_amount=total_amount,
        )

        for para in doc.paragraphs:
            # Only process paragraphs that may contain placeholders
            para_text = "".join(r.text for r in para.runs)
            if "[" not in para_text:
                continue
            _merge_runs(para)
            for old, new in replacements:
                if old in para.runs[0].text:
                    para.runs[0].text = para.runs[0].text.replace(old, new)

        # Also process tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_text = "".join(r.text for r in para.runs)
                        if "[" not in para_text:
                            continue
                        _merge_runs(para)
                        for old, new in replacements:
                            if old in para.runs[0].text:
                                para.runs[0].text = para.runs[0].text.replace(old, new)

        doc.save(output_path)
        return output_path.resolve()
